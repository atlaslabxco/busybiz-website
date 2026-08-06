# -*- coding: utf-8 -*-
"""Generate the Busy Biz mini-game build-prompt brief as a .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x16, 0x32, 0x4F)
CORAL = RGBColor(0xFF, 0x6B, 0x5E)
SLATE = RGBColor(0x5B, 0x6B, 0x7B)

doc = Document()

# base styles
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = NAVY
    p.space_after = Pt(4)
    return p

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = CORAL
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    return p

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11.5); r.font.bold = True; r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    return p

def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.italic = italic
    if color: r.font.color.rgb = color
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(10.5)
    return p

def code_block(lines):
    for ln in lines:
        p = doc.add_paragraph()
        r = p.add_run(ln)
        r.font.name = 'Consolas'; r.font.size = Pt(9.5); r.font.color.rgb = NAVY
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)

def prompt_box(title, lines):
    h3(title)
    for ln in lines:
        p = doc.add_paragraph()
        r = p.add_run(ln)
        r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(0)

# ---------------- TITLE ----------------
h1("Busy Biz Solution — Mini-Game Build Prompts")
body("Ready-to-paste prompts for building three interactive mini-games with Claude Code (Fable 5). "
     "Each game is a single, self-contained HTML file that plugs straight into the Busy Biz Solution website.",
     color=SLATE)
body("Workflow: paste one PROMPT block below into Claude Code, build the game, then hand the finished .html "
     "file back to Atlas to drop into the site's /games folder.", italic=True, color=SLATE)

# ---------------- SHARED SPEC ----------------
h2("Shared build spec (applies to ALL three games)")
body("Copy this block in FIRST, then add the specific game prompt underneath it.")
prompt_box("SHARED SPEC — paste before each game prompt:", [
    "Build a SINGLE, self-contained HTML file (inline <style> and <script>). No frameworks,",
    "no build tools, no npm. The only external resource allowed is Google Fonts.",
    "",
    "BRAND & LOOK (match the existing Busy Biz games):",
    "  - Fonts: 'Poppins' (600–800) for headings/buttons, 'Inter' for body.",
    "  - CSS variables in :root:",
    "        --navy:#16324F;  --coral:#FF6B5E;  --mist:#F7F9FB;  --slate:#5B6B7B;",
    "  - Page background: radial-gradient(circle at 50% 0%, #1f4266, #16324F).",
    "  - The game sits in a white rounded card (border-radius ~26px) centered on screen,",
    "    max-width ~460px, soft shadow. Navy text, coral accents.",
    "  - Small header row: 'busy biz' + a coral dot, then an uppercase eyebrow label",
    "    (e.g. 'RESTAURANT DEMO'). Footer line: 'Demo game by Busy Biz Solution'.",
    "",
    "BEHAVIOUR:",
    "  - Mobile-first and fully responsive. Must work with BOTH touch and mouse.",
    "  - Works standalone AND inside an <iframe>. No fixed pixel widths that overflow phones.",
    "  - End state always shows: the result, a reward CODE in a dashed coral pill,",
    "    a 'how to redeem' line, and a 'Play again' button.",
    "  - Include an OPTIONAL lead-capture step (name + email before revealing the code),",
    "    written but commented out, so it can be switched on later.",
    "",
    "REBRANDABLE (important):",
    "  - Put ALL editable content (items, prizes, questions, copy, reward codes) in clearly",
    "    labelled CONFIG arrays/objects at the top of the <script>, with comments.",
    "  - A non-coder should be able to change colours (in :root) and copy (in CONFIG) only.",
    "",
    "QUALITY:",
    "  - Large tap targets, legible sizes, no text overlap or clipping at 360px width.",
    "  - Smooth CSS/JS animations. No backend; use localStorage where persistence is needed.",
    "  - Clean, commented code.",
])

# ---------------- GAME 1 ----------------
h2("Game 1 — Restaurant: “Build Your Dish” (drag-and-drop)")
body("Concept: the player drags ingredients onto a base to build their own dish; on 'Serve' they get a "
     "discount reward on the dish they created. Hands-on, personal, and captures order intent.")
prompt_box("PROMPT — Build Your Dish:", [
    "Using the SHARED SPEC above, build a drag-and-drop restaurant mini-game called",
    "\"Build Your Dish\". Eyebrow label: 'RESTAURANT DEMO'.",
    "",
    "GAMEPLAY:",
    "  - Show a plate/base at the bottom (e.g. a burger bun or a bowl).",
    "  - Show a tray of draggable ingredients (emoji or simple shapes) above it.",
    "  - The player drags ingredients onto the base; each added ingredient stacks/",
    "    appears on the plate with a little pop animation. Support touch drag on mobile",
    "    (pointer events) AND mouse drag.",
    "  - Show a live count / mini 'flavour score' as they build.",
    "  - A 'Serve it!' button finishes the round.",
    "",
    "RESULT:",
    "  - Show the dish they built ('Your creation!') and award a discount reward on that",
    "    dish, with a code (e.g. BUILD15 = 15% off your custom dish).",
    "  - Bigger/more creative builds can unlock a better reward tier (define tiers in CONFIG).",
    "  - 'Play again' resets the plate.",
    "",
    "CONFIG at top of script:",
    "  - INGREDIENTS = [{name, emoji, points}]  (e.g. patty, cheese, lettuce, egg, sauce...)",
    "  - REWARD_TIERS = [{minItems, label, code}]",
    "  - BASE = {label, emoji}   // e.g. burger / poke bowl / pizza",
])

# ---------------- GAME 2 ----------------
h2("Game 2 — FMCG: “Swipe Your Favourite” (swipe-to-sort)")
body("Concept: a Tinder-style stack of product cards; swipe right to like, left to skip. After the stack, "
     "show a taste profile + recommended product + coupon. Fun, fast, and captures preference data.")
prompt_box("PROMPT — Swipe Your Favourite:", [
    "Using the SHARED SPEC above, build a swipe-to-sort FMCG mini-game called",
    "\"Swipe Your Favourite\". Eyebrow label: 'FMCG DEMO'.",
    "",
    "GAMEPLAY:",
    "  - Show a stack of product cards (image/emoji + name + a one-line tagline).",
    "  - Swipe RIGHT = love it, LEFT = skip. Support touch swipe AND mouse drag,",
    "    plus two big buttons (❤️ / ❌) for accessibility.",
    "  - Cards animate off-screen in the swipe direction; the next card scales up.",
    "  - Show progress (e.g. 'Card 3 of 6').",
    "",
    "RESULT:",
    "  - After the last card, compute a simple 'taste profile' from the liked tags and",
    "    show: a result title (e.g. 'You’re a Sweet-Tooth Snacker'), a RECOMMENDED product,",
    "    and a coupon code (e.g. TASTE10).",
    "  - Store the liked product names in localStorage (demo stand-in for lead data);",
    "    log them to console too.",
    "  - 'Play again' reshuffles.",
    "",
    "CONFIG at top of script:",
    "  - PRODUCTS = [{name, emoji, tagline, tags:[...]}]",
    "  - PROFILES = [{tag, title, recommend, code}]",
])

# ---------------- GAME 3 ----------------
h2("Game 3 — Event: “Tap Battle” + live leaderboard")
body("Concept: a 10-second tap-as-fast-as-you-can challenge with a local leaderboard. People queue at a booth "
     "to beat the top score. Simple, competitive, and very shareable.")
prompt_box("PROMPT — Tap Battle:", [
    "Using the SHARED SPEC above, build a fast-tap event mini-game called \"Tap Battle\".",
    "Eyebrow label: 'EVENT DEMO'.",
    "",
    "GAMEPLAY:",
    "  - A big central 'TAP!' button/target. On start, a 10-second countdown begins.",
    "  - Every tap increments a large live counter with a little bounce animation.",
    "    Support rapid touch and mouse taps (no 300ms delay; use pointer/touch events).",
    "  - When time runs out, lock input and show the final score.",
    "",
    "LEADERBOARD & RESULT:",
    "  - Ask for the player's name/initials, then save {name, score} to a localStorage",
    "    leaderboard. Show the Top 5 with the current player highlighted and their rank.",
    "  - Award a tiered reward by score (define in CONFIG), with a code.",
    "  - Include a separate 'Big screen' view toggle (URL hash #board) that shows ONLY a",
    "    large live leaderboard, for displaying on a monitor at the booth.",
    "  - 'Play again' keeps the leaderboard.",
    "",
    "CONFIG at top of script:",
    "  - GAME_SECONDS = 10",
    "  - REWARD_TIERS = [{minScore, label, code}]",
])

# ---------------- HANDOFF ----------------
h2("When the game is built — hand back to Atlas")
bullet("Send Atlas the finished .html file (or paste the full code).")
bullet("Atlas drops it into the site's /games folder and links it from the 'Try our mini-games' section.")
bullet("Tell Atlas the vertical + any colour/prize tweaks; the CONFIG blocks make that quick.")
body("Live site: https://atlaslabxco.github.io/busybiz-website/", color=SLATE)

doc.save(r"C:\Users\ericc\.openclaw\workspace\projects\busybiz-website\docs\BusyBiz-MiniGame-Build-Prompts.docx")
print("SAVED")
